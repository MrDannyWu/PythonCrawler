'''
author：DannyWu
site:   www.idannywu.com
'''

import io
import os
import json
import time
import requests
#import pandas
#import PhantomJS
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from docx.shared import Inches
from selenium import webdriver
#from get_article_contents import get_article_content


url = 'https://www.jianshu.com'

def get_one_page_article(one_page_url):
    do = webdriver.ChromeOptions()
    do.add_argument(r'user-data-dir=C:\Users\wuzhiqiang1\AppData\Local\Google\Chrome\User Data')
    browser = webdriver.Chrome("chromedriver", 0, do)
    #browser = webdriver.Chrome()
    browser.get(one_page_url)
    #articles = []
    for i in range(4):
        browser.execute_script("window.scrollTo(0, document.body.scrollHeight);") 
        time.sleep(2)    
    for j in range(20):  
        try:
            button = browser.execute_script("var a = document.getElementsByClassName('load-more'); a[0].click();")
            time.sleep(2)
        except:
            pass
    try:
        titles = browser.find_elements_by_class_name("title")
    except:
        pass
    for title in titles[:len(titles) - 2]:

        if title == 'None':
            pass
        else:
            try:
                print(title.get_attribute('href'))
                article_url = title.get_attribute('href')
                get_article_content(article_url)
            except:
                pass
 
get_one_page_article(url)


header = {
        'cookie':'signin_redirect=https%3A%2F%2Fwww.jianshu.com%2Fp%2F041e2cc55c9a; read_mode=day; default_font=font2; locale=zh-CN; _m7e_session=670535ca128f5635bdfe0c7aa3a733e3; sajssdk_2015_cross_new_user=1; sensorsdata2015jssdkcross=%7B%22distinct_id%22%3A%22165a321cb494f4-0c060c4e81d0a5-3f3c5501-1310720-165a321cb4a5c6%22%2C%22%24device_id%22%3A%22165a321cb494f4-0c060c4e81d0a5-3f3c5501-1310720-165a321cb4a5c6%22%2C%22props%22%3A%7B%22%24latest_traffic_source_type%22%3A%22%E7%9B%B4%E6%8E%A5%E6%B5%81%E9%87%8F%22%2C%22%24latest_referrer%22%3A%22%22%2C%22%24latest_referrer_host%22%3A%22%22%2C%22%24latest_search_keyword%22%3A%22%E6%9C%AA%E5%8F%96%E5%88%B0%E5%80%BC_%E7%9B%B4%E6%8E%A5%E6%89%93%E5%BC%80%22%7D%7D; Hm_lvt_0c0e9d9b1e7d617b3e6842e85b9fb068=1536040226; Hm_lpvt_0c0e9d9b1e7d617b3e6842e85b9fb068=1536040226',
        'user-agent':'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.181 Safari/537.36'
        }
def get_article_content(url):
    res = requests.get(url,headers=header)
    #print(res.text)
    soup = BeautifulSoup(res.text,'lxml')
    #print(soup)
    title = soup.select('h1.title')[0].text
    print(title)
    contents = soup.select('.show-content-free *')
    #print(contents)
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.add_heading(title,0)
    for content in contents:
        #print("###",content)
        if "data-original-src" in str(content):
            img_link = content.get('data-original-src')
            #print(img_link)
            if (img_link == None):
                pass
            else:
                print(img_link)
                try:
                    img = requests.get("https:" + img_link )
                except:
                    psaa
                if (img_link[-4:-3] == '.'):
                    file_name = "temp" + img_link[-15:]
                else:
                    file_name = "temp" + img_link[-15:] + ".jpg"
                file_path = Path(file_name)
                with open(file_name,'wb') as w:
                    w.write(img.content)
                if (file_path.exists()):
                    try:
                        doc.add_picture(file_name,width=Inches(3))
                        time.sleep(2)
                        os.remove(file_name)
                    except:
                        pass
        else:
            print(content.text)
            doc.add_paragraph(content.text)
    try:
        doc_name = str('E:\\DannyWu\\Desktop\\') + title + '.docx'
        doc_name1 = title + '.docx'
        doc.save(doc_name1)
    except:
        pass
