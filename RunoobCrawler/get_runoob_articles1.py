import requests
from bs4 import BeautifulSoup
from docx import Document
import time
from docx.oxml.ns import qn
from docx.shared import Inches
import os
from pathlib import Path

url = 'http://www.runoob.com/python3/python3-tutorial.html'
header = {
        'Cookie':'_ga=GA1.2.1658311442.1536818899; Hm_lvt_3eec0b7da6548cf07db3bc477ea905ee=1536818899,1537240172,1538023724',
        'User-Agent':'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.92 Safari/537.36'
        }

def get_article_content(article_url,header):
    web_data = requests.get(article_url,headers=header)
    status_code = web_data.status_code
    if status_code == 200:
        web_data.encoding = 'utf-8'
        try:
            soup = BeautifulSoup(web_data.text,'lxml')
            article = soup.select('.article-intro')
            for p in article:
                print(p.text)
        except:
            pass
    else:
        print('Request Error...')
        #pass


def get_article_content(url,header):
    res = requests.get(url,headers=header)
    res.encoding = 'utf-8'
    #print(res.text)
    soup = BeautifulSoup(res.text,'lxml')
    #print(soup)
    #title = soup.select('.article-intro h1')[0].text
    title = soup.select('head title')[0].text.split(' | ')[0]
    print(title)
    contents = soup.select('.article-intro *')
    #print(contents)
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.add_heading(title,0)
    for content in contents:
        #print("######",content)
        #print("###",content)
        if "img" in str(content):
            img_link = content.get('src')
            #print(img_link)
            if (img_link == None):
                pass
            else:
                print(img_link)
                try:
                    if 'http' in img_link:
                        img = requests.get(img_link )
                    elif '//' in img_link:
                        img = requests.get('http:' + img_link )
                    else:
                        img = requests.get('http://www.runoob.com' + img_link )

                except:
                    pass
                if (img_link[-4:-3] == '.'):
                    file_name = "temp" + img_link[-5:]
                else:
                    file_name = "temp" + img_link[-5:] + ".jpg"
                file_path = Path(file_name)
                with open(file_name,'wb') as w:
                    w.write(img.content)
                if (file_path.exists()):
                    try:
                        doc.add_picture(file_name,width=Inches(3))
                        time.sleep(2)
                        #os.remove(file_name)
                    except:
                        pass
        else:
            print(content.text)
            doc.add_paragraph(content.text)
    try:
        doc_name = str('E:\\DannyWu\\Desktop\\') + title + '.docx'
        doc_name1 = title + '.docx'
        doc.save(doc_name1)
    except:
        pass




def get_article_url(index_url,header):
    web_data = requests.get(index_url,headers=header)
    status_code = web_data.status_code
    if status_code == 200:
        web_data.encoding = 'utf-8'
        try:
            soup = BeautifulSoup(web_data.text,'lxml')
            article_url_list = soup.select('#leftcolumn a')
            for article_url in article_url_list:
                if article_url.get('href')[0] is '/':
                    article_url = 'http://www.runoob.com' + article_url.get('href')
                    print(article_url)
                else:
                    article_url = 'http://www.runoob.com/' + article_url.get('href')
                    print(article_url)
        except:
            pass
    else:
        print('Request Error...')



#get_article_content('http://www.runoob.com/python3/python3-basic-syntax11.html',header)
#get_article_url('http://www.runoob.com/python3/python3-basic-syntax.html',header)
get_article_content('https://www.runoob.com/js/js-intro.html',header)

